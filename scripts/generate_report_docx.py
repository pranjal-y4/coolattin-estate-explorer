"""
scripts/generate_report_docx.py

Generate a fully-formatted Word (.docx) technical implementation report
for the Coolattin Estate Records Explorer.

Usage:
    python3 scripts/generate_report_docx.py

Output:
    docs/Coolattin_Technical_Report.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

# ── Ensure project root is on sys.path ───────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT_PATH = ROOT / "docs" / "Coolattin_Technical_Report.docx"

# ── Colour palette ─────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0F, 0x17, 0x2A)   # headings
ACCENT_BLUE = RGBColor(0x1D, 0x4E, 0xD8)   # sub-headings
DEEP_GREEN  = RGBColor(0x16, 0x6A, 0x34)   # accent text
TABLE_HEAD  = RGBColor(0x1E, 0x3A, 0x5F)   # table header bg
TABLE_ALT   = RGBColor(0xF0, 0xF4, 0xFF)   # alternating row bg
CODE_BG     = RGBColor(0xF8, 0xFA, 0xFC)   # code block bg
GREY_TEXT   = RGBColor(0x47, 0x54, 0x67)   # body secondary


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def _set_para_spacing(para, before: int = 0, after: int = 6):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    if level == 1:
        run.font.color.rgb = DARK_NAVY
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(13)
        run.bold = True
    elif level == 3:
        run.font.color.rgb = DEEP_GREEN
        run.font.size = Pt(11)
        run.bold = True
    _set_para_spacing(para, before=14 if level == 1 else 10, after=4)
    return para


def add_body(doc: Document, text: str, bold_prefix: str = ""):
    para = doc.add_paragraph()
    if bold_prefix:
        run = para.add_run(bold_prefix + " ")
        run.bold = True
        run.font.color.rgb = DARK_NAVY
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    _set_para_spacing(para, before=2, after=4)
    return para


def add_bullet(doc: Document, text: str, level: int = 0, bold_prefix: str = ""):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    para = doc.add_paragraph(style=style)
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        para.add_run(" " + text).font.size = Pt(10.5)
    else:
        para.add_run(text).font.size = Pt(10.5)
    _set_para_spacing(para, before=1, after=2)
    return para


def add_code(doc: Document, code: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    _set_para_spacing(para, before=2, after=2)
    return para


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_bg(cell, TABLE_HEAD)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = TABLE_ALT if r_idx % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            _set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for r_idx, row in enumerate(table.rows):
            for c_idx, width in enumerate(col_widths):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(width)

    doc.add_paragraph()  # spacing after table
    return table


def add_divider(doc: Document):
    para = doc.add_paragraph()
    run = para.add_run("─" * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    _set_para_spacing(para, before=6, after=6)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILD
# ═══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Default paragraph font ────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_para.add_run("Coolattin Estate Records Explorer")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DARK_NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_para.add_run("Full Technical Implementation Report")
    r2.font.size = Pt(16)
    r2.font.color.rgb = ACCENT_BLUE
    r2.bold = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Masters Dissertation  ·  May {datetime.date.today().year}\n").font.size = Pt(11)
    meta.add_run("County Wicklow Estate Records · 1841–1891").font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Abstract / summary box (as a 1-cell table)
    abs_table = doc.add_table(rows=1, cols=1)
    abs_table.style = "Table Grid"
    abs_cell = abs_table.rows[0].cells[0]
    _set_cell_bg(abs_cell, RGBColor(0xEF, 0xF6, 0xFF))
    abs_p = abs_cell.paragraphs[0]
    abs_run = abs_p.add_run(
        "This report documents the complete technical implementation of the Coolattin Estate Records "
        "Explorer — a full-stack web application built as a Masters dissertation artefact. The system "
        "integrates 13,707 historical estate records (6,016 emigrations, 4,108 evictions, 5,247 tenancy "
        "records) with an interactive map, analytics dashboard, natural-language LLM Q&A, a D3.js "
        "knowledge graph visualiser, and a GraphDB SPARQL endpoint containing 143,123 RDF triples. "
        "Every feature, route, service, database table, algorithm, and design decision is documented "
        "with technical detail suitable for academic review."
    )
    abs_run.font.size = Pt(10.5)
    abs_run.font.italic = True
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Project Overview", 1)
    add_body(doc,
        "The Coolattin Estate Records Explorer is a full-stack web application providing "
        "a unified, searchable interface over historical records from the Coolattin Estate, "
        "County Wicklow, Ireland, covering the mid-nineteenth century — one of the most "
        "significant periods in Irish history, encompassing the Great Famine (1845–1852), "
        "mass emigration, and large-scale estate clearances.")

    add_heading(doc, "1.1 Historical Context", 2)
    add_body(doc,
        "The Coolattin Estate was the County Wicklow seat of the Fitzwilliam family. "
        "In the decade following the Famine, the estate oversaw:")
    add_bullet(doc, "6,016 emigration records — Fitzwilliam-sponsored passages, mainly to Quebec and Toronto", bold_prefix="Emigration:")
    add_bullet(doc, "4,108 eviction records — court sessions and estate clearances", bold_prefix="Evictions:")
    add_bullet(doc, "5,247 tenancy records — landholding relationships across the estate", bold_prefix="Tenancy:")
    add_bullet(doc, "Clearances across 122 townlands from 1847 to 1856 (Famine peak)", bold_prefix="Clearances:")
    add_bullet(doc, "Population data from 1827 to 1891 across 1,319 townlands", bold_prefix="Census:")

    add_heading(doc, "1.2 Core Objectives", 2)
    add_table(doc,
        ["Objective", "Implementation"],
        [
            ["Full-text search over 13,707 estate records", "Unified Records Search (SQLite + pandas)"],
            ["Spatial visualisation of records by townland", "Leaflet.js interactive map with GeoJSON"],
            ["Population trend analysis across time", "Census explorer + analytics dashboard"],
            ["Natural-language Q&A over historical data", "LLM pipeline (OpenRouter / Ollama + SSE)"],
            ["Knowledge graph representation of all data", "GraphDB 10.x + D3.js force-directed graph"],
            ["SQL vs SPARQL comparative prototype", "Side-by-side comparison with LLM explanation"],
            ["Reproducible academic artefact", "SQLite, deterministic ingest, no ORM"],
        ],
        col_widths=[3.0, 3.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Technology Stack", 1)

    add_heading(doc, "2.1 Backend", 2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Language",           "Python",         "3.12",  "Primary backend language"],
            ["Web framework",      "Flask",          "3.x",   "HTTP routing, SSE streaming, Jinja2 templating"],
            ["Database",           "SQLite (sqlite3)","3.x",  "Local persistence; no ORM by design"],
            ["HTTP client",        "requests",       "2.x",   "SPARQL endpoint + LLM API calls"],
            ["Data processing",    "pandas",         "2.x",   "Unified records CSV; workhouse matching"],
            ["Excel export",       "openpyxl",       "3.x",   "Census data Excel (.xlsx) export"],
            ["RDF processing",     "rdflib",         "7.x",   "In-process SPARQL against TTL file"],
            ["RDF graph database", "GraphDB",        "10.x",  "143,123-triple SPARQL endpoint (local)"],
            ["LLM (primary)",      "OpenRouter",     "API",   "Cloud LLM; 14 free models with fallback chain"],
            ["LLM (fallback)",     "Ollama",         "—",     "Local LLM server; configurable model"],
        ],
        col_widths=[1.6, 1.5, 0.9, 2.5]
    )

    add_heading(doc, "2.2 Frontend", 2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Language",    "Vanilla JavaScript", "ES2020", "All interactive logic; no framework"],
            ["Mapping",     "Leaflet.js",         "1.9",    "Interactive map with GeoJSON overlays"],
            ["Graph viz",   "D3.js",              "7",      "Force-directed knowledge graph"],
            ["Charting",    "Chart.js",           "4.x",    "Analytics bar/line/doughnut charts"],
            ["Templating",  "Jinja2",             "3.x",    "Server-side HTML rendering"],
            ["Styling",     "Custom CSS",         "—",      "Single main.css (no framework)"],
        ],
        col_widths=[1.4, 1.8, 0.9, 2.4]
    )

    add_heading(doc, "2.3 External Services", 2)
    add_table(doc,
        ["Service", "Role", "Protocol"],
        [
            ["VRTI (Virtual Record Treasury of Ireland)", "Townland metadata + census data from the Knowledge Graph", "SPARQL over HTTPS"],
            ["OpenRouter",  "Cloud LLM for SQL generation and answer rewriting", "REST (OpenAI-compatible)"],
            ["Ollama",      "Local LLM fallback when OpenRouter unavailable", "REST"],
            ["GraphDB 10.x","Local Coolattin RDF graph with 143,123 triples", "SPARQL + REST /size"],
        ],
        col_widths=[2.0, 3.0, 1.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. System Architecture", 1)

    add_heading(doc, "3.1 Layered Architecture", 2)
    add_body(doc,
        "The application follows a strict three-layer separation that ensures testability, "
        "maintainability, and academic transparency:")
    add_bullet(doc,
        "Route layer (backend/routes/*.py) — thin Flask blueprints. Each blueprint handles one URL "
        "prefix. Routes parse request parameters, call a service function, and return JSON or render "
        "a template. No business logic lives here.",
        bold_prefix="Layer 1 — Routes:")
    add_bullet(doc,
        "Service layer (backend/services/*.py) — all business logic. Services orchestrate repository "
        "calls, external API calls, caching decisions, and data transformation. The most complex service "
        "(ask_service.py) is 6,651 lines and implements the full LLM Q&A pipeline.",
        bold_prefix="Layer 2 — Services:")
    add_bullet(doc,
        "Repository layer (backend/repositories/*.py) — all SQL. Each repository module exposes typed "
        "functions that execute parameterised queries against SQLite. No raw SQL appears in services or "
        "routes, ensuring a clean data-access boundary.",
        bold_prefix="Layer 3 — Repositories:")

    add_heading(doc, "3.2 Blueprint Registration", 2)
    add_table(doc,
        ["URL Prefix", "Blueprint", "Purpose"],
        [
            ["/",               "main",        "Page routes (8 HTML templates)"],
            ["/api/census",     "census",      "Census data retrieval + export"],
            ["/api/unified",    "unified",     "Unified estate records search"],
            ["/api/ask",        "ask",         "LLM Q&A pipeline (SSE streaming)"],
            ["/api/kg",         "kg_explore",  "Knowledge graph + SPARQL comparison"],
            ["/api/map",        "map_config",  "Basemap layers + centroid coordinates"],
            ["/api/townlands",  "townlands",   "Townland reference data + KG refresh"],
            ["/api/exports",    "exports",     "Excel file download"],
        ],
        col_widths=[1.6, 1.6, 3.3]
    )

    add_heading(doc, "3.3 Key Architectural Decisions", 2)
    add_bullet(doc, "Application factory pattern: create_app() is the sole place blueprints are registered, enabling clean testing and environment-specific config.", bold_prefix="Factory pattern:")
    add_bullet(doc, "DB singleton: extensions.py exports get_db_conn() used everywhere; all connections share identical PRAGMA settings (WAL, 64 MB cache, mmap).", bold_prefix="DB singleton:")
    add_bullet(doc, "No ORM: raw sqlite3 throughout for transparency, portability, and academic reproducibility.", bold_prefix="No ORM:")
    add_bullet(doc, "SSE streaming: the Ask pipeline yields progress events so users see live updates instead of waiting for a full response.", bold_prefix="SSE streaming:")
    add_bullet(doc, "DB-first / KG-second: all slow external data (census, townlands) is cached in SQLite with a configurable TTL; the KG is only queried on cache miss or staleness.", bold_prefix="Caching strategy:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — DATABASE DESIGN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Database Design", 1)
    add_body(doc,
        "The SQLite database (coolattin.db) contains five tables, all created idempotently "
        "by extensions.py::ensure_schema(). The schema is intentionally minimal and normalised "
        "only where needed for query performance.")

    add_heading(doc, "4.1 Table: townland", 2)
    add_body(doc,
        "The canonical townland reference table. Populated from the estate GeoJSON and enriched "
        "from the VRTI Knowledge Graph. Contains 4,225 townland entries for all of Wicklow, "
        "of which 152 have full estate boundary and area data.")
    add_table(doc,
        ["Column", "Type", "Purpose"],
        [
            ["id",              "INTEGER PK",   "Auto-increment primary key"],
            ["name",            "TEXT UNIQUE",  "Canonical townland name (uppercase normalised)"],
            ["name_gaelic",     "TEXT",         "Irish-language name"],
            ["barony",          "TEXT",         "Barony (administrative division above parish)"],
            ["civil_parish",    "TEXT",         "Civil parish (indexed)"],
            ["td_id",           "TEXT",         "Estate internal identifier (TD_ID from GeoJSON)"],
            ["area_sqm",        "REAL",         "Area in square metres"],
            ["kg_uri",          "TEXT",         "VRTI Knowledge Graph URI (indexed)"],
            ["wkt_geometry",    "TEXT",         "WKT polygon boundary"],
            ["centroid_lat/lon","REAL",         "Centroid coordinates for map markers"],
            ["county",          "TEXT",         "County (indexed)"],
            ["osm_id",          "TEXT",         "OpenStreetMap identifier"],
            ["vrti_id",         "TEXT",         "VRTI internal identifier"],
            ["images_json",     "TEXT",         "JSON array of image URLs"],
        ],
        col_widths=[1.6, 1.2, 3.6]
    )

    add_heading(doc, "4.2 Table: census_record", 2)
    add_body(doc,
        "Population data per townland per year. 8,033 records across 12 survey years. "
        "Unique constraint on (townland_id, year) ensures idempotent upserts.")
    add_table(doc,
        ["Column", "Type", "Purpose"],
        [
            ["id",           "INTEGER PK",   "Auto-increment primary key"],
            ["townland_id",  "INTEGER FK",   "References townland(id)"],
            ["year",         "INTEGER",      "Survey year (1827–1891)"],
            ["male",         "INTEGER",      "Male population count"],
            ["female",       "INTEGER",      "Female population count"],
            ["total",        "INTEGER",      "Total population"],
            ["inhabited",    "INTEGER",      "Inhabited house count"],
            ["uninhabited",  "INTEGER",      "Uninhabited house count"],
            ["source",       "TEXT",         "'kg' (VRTI) or 'geojson' (estate survey)"],
            ["kg_uri",       "TEXT",         "VRTI Census KG URI"],
        ],
        col_widths=[1.6, 1.2, 3.6]
    )

    add_heading(doc, "4.3 Table: clearances_record", 2)
    add_body(doc,
        "Estate eviction clearance data per townland per year. 1,211 records covering "
        "the Famine period 1847–1856 across 122 townlands.")
    add_table(doc,
        ["Column", "Type", "Purpose"],
        [
            ["id",          "INTEGER PK",  "Auto-increment primary key"],
            ["townland_id", "INTEGER FK",  "References townland(id)"],
            ["year",        "INTEGER",     "Clearance year (1847–1856)"],
            ["count",       "INTEGER",     "Number of households/persons cleared"],
            ["source",      "TEXT",        "Data source identifier"],
        ],
        col_widths=[1.6, 1.2, 3.6]
    )

    add_heading(doc, "4.4 Table: refresh_state", 2)
    add_body(doc,
        "Tracks data freshness for cache invalidation. One row per dataset "
        "(census, townlands). Enables the DB-first / KG-second caching strategy.")

    add_heading(doc, "4.5 Table: unified_record (runtime-seeded)", 2)
    add_body(doc,
        "A denormalised view of all 13,707 estate records sourced from unified_processed.csv. "
        "Created as a temporary SQLite table at server startup by ask_service._ensure_unified_table_seeded(). "
        "Contains 60+ columns including record_id, forename, surname, townland, year, occupation, "
        "has_emigration_record (0/1), has_eviction_record (0/1), has_tenancy_record (0/1), "
        "holding_acres, family_size_estimate, rent_owed, chief_tenant_surname, is_widow, "
        "is_canada_destination, and many more.")

    add_heading(doc, "4.6 Performance Pragmas", 2)
    add_body(doc, "Applied to every SQLite connection via extensions.get_db_conn():")
    add_table(doc,
        ["PRAGMA", "Value", "Benefit"],
        [
            ["journal_mode",  "WAL",         "Write-Ahead Logging — concurrent reads during writes"],
            ["synchronous",   "NORMAL",      "Balanced durability with WAL (faster than FULL)"],
            ["cache_size",    "-65536",      "64 MB page cache — avoids repeated disk reads"],
            ["temp_store",    "2",           "In-memory temp tables — no temp file I/O"],
            ["mmap_size",     "268435456",   "256 MB memory-mapped I/O — OS-level page cache"],
            ["foreign_keys",  "ON",          "Referential integrity enforced"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — DATA INGESTION PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Data Ingestion Pipeline", 1)
    add_body(doc,
        "Data flows from three primary sources into the SQLite database through dedicated ingest jobs. "
        "All ingest jobs are idempotent (safe to re-run) and update the refresh_state table on completion.")

    add_heading(doc, "5.1 Full Ingest (backend/jobs/full_ingest.py)", 2)
    add_body(doc,
        "run_full_ingest() is the primary pipeline, combining the estate GeoJSON and VRTI KG:")

    add_bullet(doc,
        "Reads 152 townland features from frontend/static/data/townlands.json. "
        "Extracts canonical name, Irish name, estate TD_ID and GUID, area in sqm, "
        "population surveys for 1827/1839/1848/1850/1860/1868, and per-year clearances 1847–1856.",
        bold_prefix="Stage 1 — Estate GeoJSON:")
    add_bullet(doc,
        "For each townland, calls vrti_sparql.fetch_townlands() and fetch_census_records(). "
        "Retrieves WKT boundary geometry, centroid, civil parish, barony, county, OSM/OSI/VRTI IDs, "
        "images, links, and standard census 1841/1851/1861/1871/1881/1891.",
        bold_prefix="Stage 2 — VRTI KG enrichment:")
    add_bullet(doc,
        "All data is upserted via the repository layer. refresh_state table is updated with "
        "last_synced_at timestamp and record_count.",
        bold_prefix="Stage 3 — Persistence:")

    add_heading(doc, "5.2 RDF Uplift (scripts/rdf_uplift.py)", 2)
    add_body(doc,
        "Converts unified_processed.csv into RDF Turtle format and optionally loads it into GraphDB:")
    add_bullet(doc, "Each estate record becomes a co:Person instance + linked co:Event instance")
    add_bullet(doc, "Properties: schema:givenName, schema:familyName, co:townland, co:year, co:eventType, co:hasEvent")
    add_bullet(doc, "Instance URIs: ex:person_{record_id} and ex:event_{record_id}")
    add_bullet(doc, "Output: data/coolattin_sample.ttl (143,123 triples when fully loaded)")
    add_code(doc, "python3 scripts/rdf_uplift.py --import   # generate TTL + POST to GraphDB")
    add_code(doc, "python3 scripts/rdf_uplift.py --limit 500  # development sample")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — FEATURE: UNIFIED ESTATE RECORDS SEARCH
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Feature: Unified Estate Records Search", 1)
    add_body(doc,
        "The primary search interface over all 13,707 estate records. Users can filter by surname, "
        "forename, townland, year, and estate. Results include all 60+ metadata fields.")

    add_heading(doc, "6.1 Implementation", 2)
    add_body(doc,
        "unified_service.py loads unified_processed.csv once per process into a pandas DataFrame "
        "(_UNIFIED_CACHE). All subsequent searches operate entirely in memory, achieving sub-millisecond "
        "response times. Filtering uses pandas vectorised string operations:")
    add_code(doc, "mask &= df['surname'].str.upper().str.contains(surname.upper())")
    add_code(doc, "mask &= df['townland_norm'].str.contains(normalize(townland), na=False)")

    add_heading(doc, "6.2 Record Distribution", 2)
    add_table(doc,
        ["Event Type", "Records", "Percentage", "Notes"],
        [
            ["Emigration", "6,016", "43.9%", "Fitzwilliam-sponsored; mainly Quebec / Toronto"],
            ["Eviction",   "4,108", "30.0%", "Court sessions and estate clearances"],
            ["Tenancy",    "5,247", "38.3%", "Landholding relationships"],
            ["Total",      "13,707","—",     "Some records carry multiple event flags"],
        ],
        col_widths=[1.5, 1.0, 1.2, 2.8]
    )
    add_body(doc,
        "Record years span 1841–1886, concentrated in 1847–1856 (Famine clearances peak). "
        "The dataset covers 516 unique townlands and 977 distinct surnames.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — FEATURE: INTERACTIVE MAP
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Feature: Interactive Map", 1)
    add_body(doc,
        "A Leaflet.js interactive map displaying all 152 estate townland boundaries from the GeoJSON, "
        "with colour-coded choropleth overlays showing census population, eviction intensity, or "
        "emigration rates. Clicking a townland opens a detail panel with all census years and clearances.")

    add_heading(doc, "7.1 Basemap Layers", 2)
    add_table(doc,
        ["Layer", "Source", "Purpose"],
        [
            ["Standard",     "OpenStreetMap",          "Default topographic map"],
            ["Satellite",    "Esri World Imagery",      "Aerial photography overlay"],
            ["Terrain",      "OpenTopoMap",             "Topographic contour map"],
            ["Historic OSM", "historical.openstreetmap.org", "19th-century historical map (optional)"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    add_heading(doc, "7.2 Map Data", 2)
    add_bullet(doc, "GeoJSON boundaries: frontend/static/data/townlands.json (6.2 MB, 24h browser-cached)")
    add_bullet(doc, "Centroids: GET /api/centroids → {townland_name: [lat, lon]} from townland table")
    add_bullet(doc, "Choropleth colour: computed client-side from embedded population/clearances properties")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — FEATURE: CENSUS EXPLORER
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Feature: Census Explorer", 1)
    add_body(doc,
        "The census explorer shows population data across twelve survey years (1827–1891). "
        "Users filter by year, townland, or barony to observe the dramatic population decline "
        "from pre-Famine to post-Famine years directly in the data.")

    add_heading(doc, "8.1 DB-First / KG-Second Caching", 2)
    add_body(doc, "The canonical caching strategy used throughout the application:")
    add_table(doc,
        ["Condition", "Action", "Response cache_status"],
        [
            ["DB hit, data fresh (within TTL)",  "Query SQLite, return immediately",             "\"hit\""],
            ["DB hit, data stale (TTL expired)",  "Return DB copy + queue background KG refresh", "\"stale_refresh\""],
            ["DB miss",                           "Query VRTI SPARQL → persist → generate Excel", "\"miss\""],
        ],
        col_widths=[2.3, 2.7, 1.5]
    )
    add_body(doc,
        "TTL: 7 days in development, 1 day in production (CENSUS_STALE_AFTER_DAYS config).")

    add_heading(doc, "8.2 Census Data Coverage", 2)
    add_table(doc,
        ["Year", "Source", "Notes"],
        [
            ["1827, 1839, 1848, 1850, 1860, 1868", "Estate GeoJSON surveys",     "Total population only (no gender split)"],
            ["1841, 1851, 1861, 1871, 1881, 1891",  "VRTI Knowledge Graph (KG)", "Male + female + inhabited/uninhabited houses"],
        ],
        col_widths=[2.5, 2.0, 2.0]
    )
    add_body(doc, "Total: 8,033 census records across 12 years and 1,319 townlands.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — FEATURE: ANALYTICS DASHBOARD
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Feature: Analytics Dashboard", 1)
    add_body(doc,
        "A pluggable analytics dashboard showing KPIs and Chart.js visualisations for each dataset. "
        "The framework is designed for extensibility: adding a new dataset requires only creating "
        "a module file with a MODULE object — no registration step needed.")

    add_heading(doc, "9.1 AnalyticsModule Protocol", 2)
    add_code(doc, "class AnalyticsModule(Protocol):")
    add_code(doc, "    dataset_id:   str")
    add_code(doc, "    dataset_name: str")
    add_code(doc, "    description:  str")
    add_code(doc, "    def compute(self) -> AnalyticsResult: ...")
    add_body(doc,
        "analytics/registry.py auto-discovers all .py files in the analytics/ directory "
        "and collects those defining a MODULE attribute.")

    add_heading(doc, "9.2 Current Analytics Modules", 2)
    add_table(doc,
        ["Module", "Dataset ID", "Key KPIs / Charts"],
        [
            ["emigrations.py",  "emigration_stats",  "Total emigrants, year trend, top townlands, Canada vs other destinations"],
            ["evictions.py",    "evictions_stats",   "Total evictions, per-year trend, clearances breakdown, geographic distribution"],
            ["tenancies.py",    "tenancy_stats",     "Tenancy count, family size distribution, holding acres histogram"],
            ["townland_geo.py", "townland_geography","Townland count, area distribution, parish breakdown"],
            ["unified.py",      "unified_overview",  "Total records, surname frequency top-10, field coverage metrics"],
        ],
        col_widths=[1.5, 1.7, 3.3]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — FEATURE: NATURAL-LANGUAGE ASK (LLM Q&A)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Feature: Natural-Language Ask (LLM Q&A)", 1)
    add_body(doc,
        "The Ask page accepts a free-text historical research question, processes it through an "
        "11-stage pipeline, and returns a precise, data-backed answer. The pipeline combines rule-based "
        "SQL template matching (fast path, no LLM) with LLM-generated SQL (slow path), VRTI and GraphDB "
        "enrichment, natural-language rewriting, and PDF report generation.")

    add_heading(doc, "10.1 Pipeline Stages (SSE Events)", 2)
    add_table(doc,
        ["Stage", "Event Type", "Description"],
        [
            ["1",  "initializing",          "Pipeline starts"],
            ["2",  "loading_context",        "Seed unified_record + heritage tables from CSV/GeoJSON"],
            ["3",  "running_template_match", "Keyword-scored SQL template search (100+ templates)"],
            ["4",  "townland_resolution",    "Exact → fuzzy → alias → 'did you mean?' matching"],
            ["5",  "contacting_llm",         "Triggered only if no template matched"],
            ["6",  "framing_query",          "LLM generating parameterised SQL"],
            ["7",  "querying_database",      "Execute validated SQL against SQLite"],
            ["8",  "querying_vrti_graph",    "VRTI SPARQL enrichment (parallel with DB query)"],
            ["9",  "querying_graphdb",       "GraphDB SPARQL (if enabled) for SQL vs SPARQL comparison"],
            ["10", "preparing_output",       "LLM rewrite to natural language + PDF generation"],
            ["11", "complete",               "Final result JSON emitted"],
        ],
        col_widths=[0.5, 2.0, 4.0]
    )

    add_heading(doc, "10.2 Template Matching (Fast Path)", 2)
    add_body(doc,
        "The most impactful performance optimisation: 100+ pre-verified SQL templates matched by "
        "keyword scoring using difflib.SequenceMatcher. Response time without LLM: < 100ms. "
        "Response time with LLM: 3–15 seconds. Template matching succeeds for approximately 60% "
        "of common research questions.")
    add_body(doc,
        "15 Verified Analysis Templates handle high-confidence research queries and include "
        "chart_hint metadata (bar, line, doughnut) for frontend visualisation: "
        "tenant_land_gender_average, widows_with_children_proportion, "
        "emigration_population_townland_trend, clearances_eviction_comparison, etc.")

    add_heading(doc, "10.3 Townland Resolution", 2)
    add_table(doc,
        ["Step", "Method", "Threshold"],
        [
            ["1 — Exact match",  "UPPER(name) = UPPER(input) in SQLite townland table",       "—"],
            ["2 — Fuzzy match",  "difflib.SequenceMatcher ratio",                              "≥ 0.72"],
            ["3 — Alias lookup", "townland_aliases.json maps variant spellings to canonical",   "—"],
        ],
        col_widths=[1.8, 3.5, 1.2]
    )
    add_body(doc,
        "If confidence < 0.85, a 'did you mean?' suggestion is added. "
        "If no match found, alternative townland suggestions are returned.")

    add_heading(doc, "10.4 LLM SQL Generation", 2)
    add_body(doc, "The SQL generation prompt includes:")
    add_bullet(doc, "Live SQLite schema with row counts and sampled categorical values")
    add_bullet(doc, "Previously approved queries from ask_query_memory (semantic similarity ranking)")
    add_bullet(doc, "Mandatory rules: COUNT(DISTINCT record_id), NEVER GROUP_CONCAT, LIMIT 50 for person lists")
    add_bullet(doc, "FORBIDDEN properties list preventing LLM from confusing SQLite columns with RDF predicates")
    add_body(doc, "LLM provider fallback chain (ASK_LLM_PROVIDER=auto):")
    add_bullet(doc, "1. OpenRouter (primary) — preferred model (gpt-oss-20b:free)")
    add_bullet(doc, "2. 13 additional free OpenRouter models tried in sequence on failure")
    add_bullet(doc, "3. Ollama local instance (final fallback)")

    add_heading(doc, "10.5 SQL Safety Guardrail", 2)
    add_body(doc,
        "Before any LLM-generated SQL is executed, a strict regex rejects all write operations:")
    add_code(doc, r"FORBIDDEN = re.compile(r'\b(INSERT|UPDATE|DELETE|DROP|ALTER|CREATE|REPLACE")
    add_code(doc, r"                        |TRUNCATE|PRAGMA(?!table_info)|ATTACH|DETACH)\b', re.I)")
    add_body(doc, "Only SELECT statements and CTEs (WITH … SELECT) pass this guardrail.")

    add_heading(doc, "10.6 LLM Answer Rewriting", 2)
    add_body(doc,
        "_generate_rephrased_answer() rewrites the raw SQL result into 1–3 sentences of plain English. "
        "The prompt enforces: use only supplied data, keep all numbers identical to the actual result, "
        "for results with >10 rows state the count and give 1–2 examples (do not enumerate all). "
        "Individual cell values are pre-truncated to 300 characters before being sent to the LLM "
        "(_build_llm_data_context._truncate_row()), preventing prompt inflation.")

    add_heading(doc, "10.7 Query Memory (Learning from Feedback)", 2)
    add_body(doc,
        "Users rate answers with thumbs-up/down via POST /api/ask/feedback. Approved queries are "
        "stored in ask_query_memory with SQL text, result summary, and confidence score. "
        "_find_similar_approved_queries() uses trigram + SequenceMatcher hybrid scoring to find "
        "previously approved queries matching the current question. These are injected into the "
        "LLM prompt as examples, improving accuracy over time.")

    add_heading(doc, "10.8 GraphDB SPARQL Comparison (D8 Prototype)", 2)
    add_body(doc,
        "When GRAPHDB_ENABLED=true, the pipeline also generates and executes the equivalent SPARQL "
        "query against the local GraphDB instance. SPARQL generation uses three steps:")
    add_bullet(doc, "Rule-based template matching (_match_sparql_template()) — 15 common patterns with verified SPARQL", bold_prefix="Step 1:")
    add_bullet(doc, "LLM generation with ontology schema + FORBIDDEN properties list", bold_prefix="Step 2:")
    add_bullet(doc, "Post-validation (_sparql_uses_forbidden_props()) — fallback if banned properties detected", bold_prefix="Step 3:")
    add_body(doc,
        "If SQLite and GraphDB return different values, _explain_result_mismatch() calls the LLM "
        "with both queries and results to generate an academic explanation of the discrepancy "
        "(closed-world vs open-world semantics, NULL handling, schema modelling differences).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — FEATURE: KNOWLEDGE GRAPH VISUALISER
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. Feature: Knowledge Graph Visualiser", 1)
    add_body(doc,
        "An interactive D3.js v7 force-directed graph visualising the entire Coolattin Estate "
        "database as a knowledge graph. The graph represents all 13,707 estate records aggregated "
        "across 516 townlands, grouped under 20 civil parishes, connected to 3 event hubs "
        "(Emigration, Eviction, Tenancy) and 2 data hubs (Census Records, Clearances).")

    add_heading(doc, "11.1 Graph Statistics", 2)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Total nodes",             "541"],
            ["Total edges",             "1,254"],
            ["Event hub nodes",         "3 (Emigration, Eviction, Tenancy)"],
            ["Data hub nodes",          "2 (Census Records, Clearances)"],
            ["Parish nodes",            "20"],
            ["Townland nodes",          "516"],
            ["parish_townland edges",   "178"],
            ["townland_event edges",    "820"],
            ["townland_census edges",   "147"],
            ["townland_clearances edges","109"],
        ],
        col_widths=[3.0, 3.5]
    )

    add_heading(doc, "11.2 Node Types and Visual Encoding", 2)
    add_table(doc,
        ["Node Type", "Count", "Colour", "Size", "Fixed Position"],
        [
            ["EventHub — Emigration",  "1",   "#16a34a (green)",  "22px", "Top-left (12%, 18%)"],
            ["EventHub — Eviction",    "1",   "#dc2626 (red)",    "22px", "Top-right (88%, 18%)"],
            ["EventHub — Tenancy",     "1",   "#2563eb (blue)",   "22px", "Bottom-centre (50%, 88%)"],
            ["DataHub — Census",       "1",   "#0891b2 (teal)",   "18px", "Bottom-left (12%, 82%)"],
            ["DataHub — Clearances",   "1",   "#9333ea (purple)", "18px", "Bottom-right (88%, 82%)"],
            ["Parish",                 "20",  "#7c3aed (violet)", "16px", "Float (D3 simulation)"],
            ["Townland — emigration",  "N",   "#15803d (dark green)","8–22px","Float"],
            ["Townland — eviction",    "N",   "#b91c1c (dark red)","8–22px","Float"],
            ["Townland — tenancy",     "N",   "#1d4ed8 (dark blue)","8–22px","Float"],
        ],
        col_widths=[2.0, 0.7, 1.8, 0.8, 1.8]
    )
    add_body(doc,
        "Townland node size is scaled by total record count: size = min(8 + total_records / 30, 22). "
        "Carnew (550 records) renders at maximum size 22; small townlands (≤ 10 records) at size 8.")

    add_heading(doc, "11.3 D3 Force Simulation Parameters", 2)
    add_table(doc,
        ["Force", "Parameter", "Value", "Effect"],
        [
            ["forceLink",     "distance — parish_townland",    "55px",    "Parishes cluster close to their townlands"],
            ["forceLink",     "distance — townland_event",     "110px",   "Townlands pulled toward event hubs"],
            ["forceLink",     "strength — parish_townland",    "0.55",    "Strong grouping of parish clusters"],
            ["forceLink",     "strength — townland_event",     "0.25",    "Moderate pull toward hubs"],
            ["forceManyBody", "strength — EventHub",           "-900",    "Strong hub repulsion keeps hubs apart"],
            ["forceManyBody", "strength — Parish",             "-350",    "Moderate parish separation"],
            ["forceManyBody", "strength — Townland",           "-45",     "Mild townland spread"],
            ["alphaDecay",    "—",                             "0.022",   "Longer settle time for cleaner layout"],
        ],
        col_widths=[1.5, 2.2, 0.9, 2.0]
    )

    add_heading(doc, "11.4 Backend Graph Builder", 2)
    add_body(doc,
        "kg_service._build_graph_inner() builds the graph from four SQLite queries and caches "
        "the result in a process-level singleton (_GRAPH_CACHE):")
    add_bullet(doc, "Aggregated stats per townland: GROUP BY townland with SUM(has_*_record) counts")
    add_bullet(doc, "Parish mapping from unified_record.parish (most common clean value per townland)")
    add_bullet(doc, "Fallback enrichment from townland reference table (UPPER key lookup)")
    add_bullet(doc, "Census/clearances membership: SELECT DISTINCT townland_id from each table")

    add_heading(doc, "11.5 On-Demand Person Drill-Down", 2)
    add_body(doc,
        "Clicking any Townland node triggers an async fetch to GET /api/kg/townland/<name>. "
        "kg_service.get_townland_persons() returns up to 50 persons with name, year, event type, "
        "and occupation. The detail panel shows aggregated stats, parish/barony/county, and "
        "a scrollable person list with event-type colour badges.")

    add_heading(doc, "11.6 Client-Side Search", 2)
    add_body(doc,
        "The search box filters nodes client-side across label, civil_parish, barony, county, "
        "dominant_event, and type fields. Matching nodes are highlighted (opacity=1, gold stroke, "
        "slightly enlarged); non-matching nodes fade to opacity 0.06. Connected edges of matching "
        "nodes increase to stroke-opacity 0.55; others drop to 0.03.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 12 — FEATURE: GRAPHDB SPARQL INTEGRATION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Feature: GraphDB SPARQL Integration", 1)

    add_heading(doc, "12.1 RDF Ontology", 2)
    add_body(doc, "The Coolattin estate data uses a custom ontology with the following namespaces:")
    add_table(doc,
        ["Prefix", "Namespace URI", "Purpose"],
        [
            ["co:",     "https://coolattin.ie/ontology#", "Coolattin domain ontology (Person, Event, properties)"],
            ["ex:",     "https://coolattin.ie/resource/",  "Instance URIs (ex:person_1234, ex:event_1234)"],
            ["schema:", "https://schema.org/",             "Standard vocabulary (givenName, familyName)"],
            ["xsd:",    "http://www.w3.org/2001/XMLSchema#","Data types (integer, string)"],
        ],
        col_widths=[0.9, 2.8, 2.8]
    )
    add_body(doc, "Classes: co:Person, co:Event")
    add_body(doc, "Key properties: schema:givenName, schema:familyName, co:townland, co:estate, co:hasEvent, co:eventType, co:year")

    add_heading(doc, "12.2 GraphDB Client (graphdb_sparql.py)", 2)
    add_body(doc,
        "The sole module that communicates with GraphDB. All queries use POST (not GET) per "
        "SPARQL 1.1 Protocol §2.1.3 — this instance stalls indefinitely on GET requests:")
    add_code(doc, "resp = requests.post(endpoint,")
    add_code(doc, "    data={'query': full_query},")
    add_code(doc, "    headers={'Accept': 'application/sparql-results+json',")
    add_code(doc, "             'Content-Type': 'application/x-www-form-urlencoded'})")
    add_body(doc,
        "probe() and triple_count() use the GraphDB REST /size endpoint (returns plain-text integer "
        "instantly) rather than a SPARQL query. All calls wrap exceptions and return ([], []) on "
        "failure — the application degrades gracefully when GraphDB is unavailable.")

    add_heading(doc, "12.3 Comparison Scenarios", 2)
    add_table(doc,
        ["Scenario ID", "Description", "Key Semantic Difference"],
        [
            ["emigration_count_by_townland", "Emigrants grouped by townland", "SQL NULLs excluded by WHERE clause; SPARQL naturally excludes unbound ?townland"],
            ["eviction_count_by_year",       "Evictions per year",           "4 NULL year values cause SQL rowcount difference unless explicitly excluded"],
            ["surname_frequency",            "Top 10 surnames",             "Usually identical — cleanest parity example; validates both systems agree"],
            ["person_event_detail",          "Person + event with join",    "SPARQL graph walk (co:hasEvent) vs SQL CASE WHEN event flag columns"],
        ],
        col_widths=[1.7, 2.0, 2.8]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 13 — ADDITIONAL FEATURES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13. Additional Features", 1)

    add_heading(doc, "13.1 PDF Export", 2)
    add_body(doc,
        "After every Ask pipeline run, a PDF report is generated in hand-written PDF 1.4 format "
        "(no reportlab/fpdf dependency). The writer constructs the binary PDF directly: catalog, "
        "pages tree, font object (Helvetica), content stream with BT…ET text blocks, XRef table, "
        "and trailer. Content includes: question, SQL query, data table, LLM-rewritten answer, "
        "VRTI enrichment, townland context, and LLM metadata.")

    add_heading(doc, "13.2 Excel Export", 2)
    add_body(doc,
        "Census data is exported as a structured .xlsx workbook using openpyxl with two sheets: "
        "a data sheet (townland, year, male, female, total, inhabited, uninhabited) and a metadata "
        "sheet (source, generated_at, query scope, record count). "
        "Files are named census_wicklow_<scope>_YYYYMMDD_HHMMSS.xlsx.")

    add_heading(doc, "13.3 Workhouse Fuzzy Matching", 2)
    add_body(doc,
        "Each unified estate record can be cross-referenced against workhouse admission/discharge "
        "registers using a 4-stage place-first matching algorithm:")
    add_table(doc,
        ["Stage", "Operation", "Details"],
        [
            ["1 — Place filter",    "electoral_division match",         "Filter workhouse records where electoral division matches estate townland/parish"],
            ["2 — Date window",     "Year proximity filter",            "Retain candidates within ±1 year of estate record year"],
            ["3 — Name scoring",    "difflib.SequenceMatcher",          "Score concatenated forename + surname strings"],
            ["4 — Occupation bonus","+0.05 if keywords match",         "Shared keywords: labourer, farmer, servant, widow, etc."],
        ],
        col_widths=[1.5, 1.8, 3.2]
    )
    add_table(doc,
        ["Confidence Band", "Criteria"],
        [
            ["High",   "Place match AND date match AND score ≥ 0.80"],
            ["Medium", "(Place OR date match) AND score ≥ 0.60"],
            ["Low",    "Score ≥ 0.60, no place or date match"],
        ],
        col_widths=[1.5, 5.0]
    )

    add_heading(doc, "13.4 Historic Landscape / Heritage Layer", 2)
    add_body(doc,
        "A dedicated Leaflet.js map view showing archaeological features across the estate area. "
        "Overlays: holy wells (holywells_wicklow.geojson), ring forts and souterrains "
        "(asi_wicklow.geojson from the Archaeological Survey of Ireland). "
        "These GeoJSON files are also seeded into temporary SQLite tables (holy_wells, ring_forts) "
        "at Ask startup, enabling spatial SQL queries like distance_km() radius searches.")

    add_heading(doc, "13.5 Internationalisation (English / Irish)", 2)
    add_body(doc,
        "A lightweight client-side i18n framework (i18n.js) supports English and Irish (Gaeilge). "
        "HTML elements with data-i18n='key' are auto-translated on page load. Language toggle in "
        "the navigation bar switches between EN/GA without a page reload. "
        "Translation coverage: all navigation, headings, CTA buttons, form labels, UI strings.")

    add_heading(doc, "13.6 Townland Autocomplete", 2)
    add_body(doc,
        "GET /api/ask/townland-suggest?q=<query> returns fuzzy-matched townland suggestions using "
        "difflib.SequenceMatcher against the full townland list. Results are ranked by similarity "
        "ratio and returned as {name, confidence} pairs. An alias map (townland_aliases.json) "
        "maps ~200 spelling variants to canonical names.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 14 — SECURITY & PERFORMANCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14. Security and Performance", 1)

    add_heading(doc, "14.1 Security Measures", 2)
    add_table(doc,
        ["Threat", "Mitigation", "Implementation Location"],
        [
            ["SQL injection",         "Parameterised queries throughout; FORBIDDEN_SQL regex blocks all write ops", "All repository files + ask_service.py"],
            ["XSS (stored/reflected)","escapeHtml() on all user-supplied values before innerHTML assignment", "All frontend JS files"],
            ["CSRF",                  "Flask SESSION_COOKIE_SECURE + SECRET_KEY", "config.py + create_app.py"],
            ["LLM prompt injection",  "SQL guardrail + schema constraints; LLM cannot read secrets from env", "ask_service.py"],
            ["API key exposure",      ".env.local gitignored; .env.example shows safe placeholders only", ".gitignore"],
        ],
        col_widths=[1.5, 2.8, 2.2]
    )

    add_heading(doc, "14.2 Performance Optimisations", 2)
    add_table(doc,
        ["Optimisation", "Mechanism", "Benefit"],
        [
            ["WAL mode",            "PRAGMA journal_mode=WAL",                       "Concurrent reads during write operations"],
            ["64 MB page cache",    "PRAGMA cache_size=-65536",                      "Avoids repeated disk reads for hot pages"],
            ["256 MB mmap",         "PRAGMA mmap_size=268435456",                    "OS-level page cache bypass"],
            ["Covering indexes",    "(townland_id, year) composite on census+clearances", "Avoids table scans on common join patterns"],
            ["Process-level cache", "pandas DataFrame for unified_processed.csv",    "Sub-millisecond record search after warmup"],
            ["KG topology cache",   "_GRAPH_CACHE singleton in kg_service",          "D3.js graph built once per server process"],
            ["Template matching",   "100+ SQL templates before LLM call",            "~60% of queries answered in < 100ms"],
            ["Browser cache",       "24h Cache-Control on static files",             "6.2 MB townlands.json served once per day"],
            ["Query memory",        "ask_query_memory table with similarity ranking","Previously approved queries reused without LLM"],
        ],
        col_widths=[1.8, 2.4, 2.3]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 15 — COMPLETE API REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "15. Complete API Reference", 1)

    add_heading(doc, "15.1 Page Routes", 2)
    add_table(doc,
        ["Method", "Path", "Template"],
        [
            ["GET", "/",                  "index.html — Home page with hero, records overview"],
            ["GET", "/about",             "about.html — Project and estate context"],
            ["GET", "/analytics",         "analytics.html — KPI + chart dashboard"],
            ["GET", "/census",            "census.html — Census explorer with map"],
            ["GET", "/info",              "info.html — Estate and Famine historical context"],
            ["GET", "/ask",               "ask.html — Natural-language LLM Q&A interface"],
            ["GET", "/heritage",          "heritage.html — Historic landscape overlay map"],
            ["GET", "/explore-knowledge", "kg_explore.html — Knowledge graph visualiser"],
        ],
        col_widths=[0.7, 2.0, 3.8]
    )

    add_heading(doc, "15.2 Census API (/api/census)", 2)
    add_table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET",  "/",                  "Paginated records (year, townland, barony, page, limit filters)"],
            ["GET",  "/townlands",          "List townlands with census data"],
            ["GET",  "/summary",            "Aggregate stats by year"],
            ["GET",  "/townland?name=",     "Full population history for one townland (all years)"],
            ["POST", "/refresh",            "Force VRTI KG re-ingestion (returns 202 Accepted)"],
            ["GET",  "/export/latest",      "Metadata for most recent .xlsx export"],
            ["POST", "/export/regenerate",  "Regenerate Excel from current DB (no KG call)"],
        ],
        col_widths=[0.7, 2.0, 3.8]
    )

    add_heading(doc, "15.3 Ask API (/api/ask)", 2)
    add_table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["POST", "/query",              "SSE pipeline stream; body: {question, townland_hint?, show_sql?, force_llm?}"],
            ["POST", "/feedback",           "Record thumbs-up/down + metadata for query memory"],
            ["GET",  "/llm-status",         "LLM provider health check"],
            ["GET",  "/townland-suggest?q=","Townland name autocomplete"],
            ["GET",  "/pdf/<filename>",     "Download generated PDF report"],
        ],
        col_widths=[0.7, 2.0, 3.8]
    )

    add_heading(doc, "15.4 Knowledge Graph API (/api/kg)", 2)
    add_table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET",  "/graph",              "Graph nodes + edges (541 nodes, 1,254 edges) for D3.js"],
            ["GET",  "/scenarios",          "Canned SQL vs SPARQL comparison scenarios"],
            ["POST", "/compare",            "Execute scenario or custom SQL/SPARQL; returns both result sets"],
            ["POST", "/explain-mismatch",   "LLM explanation of why SQL and SPARQL results differ"],
            ["GET",  "/graphdb-status",     "Live GraphDB connectivity; returns triple_count"],
            ["GET",  "/townland/<name>",    "Person records for a specific townland (drill-down)"],
            ["GET",  "/rdf-status",         "TTL file presence, size, and in-process triple count"],
        ],
        col_widths=[0.7, 2.0, 3.8]
    )

    add_heading(doc, "15.5 Other APIs", 2)
    add_table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET", "/api/unified/records",       "Search records (surname, forename, townland, year, estate, limit)"],
            ["GET", "/api/unified/stats",          "Total records and field coverage statistics"],
            ["GET", "/api/unified/surname-suggest","Surname autocomplete"],
            ["GET", "/api/townlands/",             "All townlands (optional ?county= filter)"],
            ["POST","/api/townlands/refresh",      "Force VRTI refresh of townland reference data"],
            ["GET", "/api/map/layers",             "Basemap tile layer definitions"],
            ["GET", "/api/centroids",              "Townland centroid coordinates {name: [lat, lon]}"],
            ["GET", "/api/exports/census/download","Download latest census .xlsx file"],
        ],
        col_widths=[0.7, 2.3, 3.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 16 — CODEBASE METRICS & SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "16. Codebase Metrics and Summary", 1)

    add_heading(doc, "16.1 Code Scale", 2)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Python files",                     "75"],
            ["JavaScript files (frontend/static)","11"],
            ["HTML Jinja2 templates",             "9"],
            ["Lines in ask_service.py",           "6,651"],
            ["Functions in ask_service.py",       "154"],
            ["Total Python LOC (approx.)",        "~12,000"],
            ["Total JavaScript LOC (approx.)",    "~4,000"],
            ["Flask blueprints",                  "8"],
            ["API endpoints",                     "30+"],
            ["SQL templates (Ask)",               "100+"],
            ["Verified analysis templates",       "15"],
            ["SPARQL comparison scenarios",       "4"],
            ["Analytics modules",                 "5"],
        ],
        col_widths=[3.0, 3.5]
    )

    add_heading(doc, "16.2 Data Statistics", 2)
    add_table(doc,
        ["Dataset", "Records", "Notes"],
        [
            ["Estate records (unified_record)", "13,707",   "Source: unified_processed.csv (checked into git)"],
            ["Townland references",             "4,225",    "Source: VRTI KG (all Wicklow townlands)"],
            ["Census records",                  "8,033",    "12 years × 1,319 townlands"],
            ["Clearances records",              "1,211",    "10 years (1847–1856) × 122 townlands"],
            ["GraphDB RDF triples",             "143,123",  "Full dataset in Coolattin co: ontology"],
            ["Unique townlands (estate)",       "516",      "From unified_record table"],
            ["Unique surnames",                 "977",      "From unified_record table"],
            ["Civil parishes (reference)",      "22",       "From townland reference table"],
            ["Graph nodes (KG visualiser)",     "541",      "20 parishes + 516 townlands + 5 hubs"],
            ["Graph edges (KG visualiser)",     "1,254",    "parish_townland + event + census + clearances"],
        ],
        col_widths=[2.5, 1.0, 3.0]
    )

    add_heading(doc, "16.3 Feature Summary", 2)
    add_table(doc,
        ["#", "Feature", "Key Technologies"],
        [
            ["1",  "Unified estate records search",              "pandas, SQLite, Jinja2"],
            ["2",  "Interactive choropleth map",                 "Leaflet.js, GeoJSON, SQLite"],
            ["3",  "Census explorer (12 years)",                 "SQLite, VRTI SPARQL, openpyxl"],
            ["4",  "Pluggable analytics dashboard",              "Chart.js, Protocol-based modules"],
            ["5",  "Historic heritage landscape overlay",        "Leaflet.js, GeoJSON overlays"],
            ["6",  "Natural-language Q&A with LLM",             "OpenRouter, Ollama, SSE streaming"],
            ["7",  "SQL template matching (fast path, no LLM)", "difflib, 100+ templates"],
            ["8",  "VRTI Knowledge Graph enrichment",           "SPARQL, external endpoint"],
            ["9",  "GraphDB SPARQL integration",                "GraphDB 10.x, co: ontology, POST"],
            ["10", "SQL vs SPARQL comparison",                  "rdflib, LLM mismatch analysis"],
            ["11", "D3.js knowledge graph visualiser",          "D3.js v7, force simulation"],
            ["12", "Townland person drill-down",                "SQLite, async fetch"],
            ["13", "Hand-written PDF export (no library)",      "PDF 1.4 binary writer"],
            ["14", "Excel export with metadata sheet",          "openpyxl"],
            ["15", "Workhouse fuzzy matching",                  "pandas, difflib SequenceMatcher"],
            ["16", "Feedback / query memory learning",          "SQLite, semantic similarity"],
            ["17", "Townland autocomplete + alias resolution",  "fuzzy matching, alias JSON"],
            ["18", "Internationalisation (English / Irish)",    "Vanilla JS i18n framework"],
            ["19", "RDF uplift (CSV → TTL → GraphDB)",         "rdflib, custom co: ontology"],
            ["20", "DB-first / KG-second caching framework",   "SQLite TTL, SPARQL on miss"],
        ],
        col_widths=[0.4, 2.8, 2.8]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # FOOTER / CLOSING
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_divider(doc)

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run(
        "Coolattin Estate Records Explorer  ·  Masters Dissertation  ·  "
        f"Generated {datetime.date.today().strftime('%d %B %Y')}"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = GREY_TEXT
    r.italic = True

    # ── Save ─────────────────────────────────────────────────────────────────
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Report saved → {OUT_PATH}")


if __name__ == "__main__":
    build()
