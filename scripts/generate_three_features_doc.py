from __future__ import annotations

import datetime
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = ROOT / "docs" / "Three_Feature_Implementation.docx"

DARK_NAVY   = RGBColor(0x0F, 0x17, 0x2A)
ACCENT_BLUE = RGBColor(0x1D, 0x4E, 0xD8)
DEEP_GREEN  = RGBColor(0x16, 0x6A, 0x34)
TABLE_HEAD  = RGBColor(0x1E, 0x3A, 0x5F)
TABLE_ALT   = RGBColor(0xF0, 0xF4, 0xFF)
CODE_BG     = RGBColor(0xF8, 0xFA, 0xFC)
GREY_TEXT   = RGBColor(0x47, 0x54, 0x67)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = DARK_NAVY
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(13)
        run.font.bold = True
    elif level == 3:
        run.font.color.rgb = DEEP_GREEN
        run.font.size = Pt(11)
        run.font.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5FB")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], TABLE_HEAD)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        bg = TABLE_ALT if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for ci, val in enumerate(row_data):
            row_cells[ci].text = str(val)
            set_cell_bg(row_cells[ci], bg)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Coolattin Estate Records Explorer")
tr.font.size = Pt(22)
tr.font.bold = True
tr.font.color.rgb = DARK_NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Three Feature Implementation Report")
sr.font.size = Pt(14)
sr.font.color.rgb = ACCENT_BLUE

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run(f"Masters Dissertation  ·  {datetime.date.today().strftime('%B %Y')}")
dr.font.size = Pt(11)
dr.font.color.rgb = GREY_TEXT

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

add_heading(doc, "Overview", level=1)
add_body(doc,
    "This document describes the design and implementation of three technical "
    "features developed for the Coolattin Estate Records Explorer as part of a "
    "Masters dissertation project. The application is a Python/Flask web application "
    "that integrates mid-19th century Irish estate records — tenancy, eviction, emigration, "
    "and census data — into a unified, searchable interface backed by SQLite and the VRTI "
    "Knowledge Graph SPARQL endpoint."
)
add_body(doc,
    "The three features documented here address distinct challenges: cross-dataset "
    "person matching (Workhouse Matching), database and in-process performance under "
    "concurrent load (Application Performance), and side-by-side comparison of "
    "relational and RDF query paradigms over the same underlying data (RDF/KG Comparison)."
)

add_table(doc,
    headers=["Feature", "Primary File(s)", "Key Technology"],
    rows=[
        ["Workhouse Matching",    "backend/services/workhouse_service.py",               "difflib.SequenceMatcher · pandas · in-process index"],
        ["Application Performance", "extensions.py · backend/services/ask_service.py",  "SQLite PRAGMAs · composite indexes · thread-safe caches"],
        ["RDF/KG Comparison",     "backend/services/kg_service.py\nbackend/routes/kg_explore.py", "rdflib · GraphDB HTTP · OpenRouter LLM · D3.js"],
    ],
    col_widths=[1.8, 2.8, 2.5],
)

add_divider(doc)

add_heading(doc, "Feature 1: Workhouse Matching", level=1)

add_heading(doc, "1.1  Problem Statement", level=2)
add_body(doc,
    "The Shillelagh Union Workhouse register (workhouse_data_final.xlsx) lists paupers "
    "admitted or born in the workhouse during the Famine period. These records use "
    "free-text name fields, inconsistent date formats, and geographic designations "
    "(electoral divisions) that overlap only partially with the townland and parish "
    "fields in the main unified estate records. The challenge is to produce ranked, "
    "confidence-banded matches between the ~13,000 unified estate records and the "
    "workhouse register without a shared primary key."
)

add_heading(doc, "1.2  Data Sources", level=2)
add_body(doc,
    "The workhouse register is held in two Excel sheets within workhouse_data_final.xlsx:"
)
add_table(doc,
    headers=["Sheet", "Name field", "Key attributes"],
    rows=[
        ["1-127",     "Pauper Name (Surname Forename)",            "Register number only; no electoral division or dates"],
        ["from 128",  "Names and Surnames of Paupers (Surname Forename)",
                       "Electoral division · Sex · Age · Employment · Admission date · Departure date"],
    ],
    col_widths=[1.0, 2.2, 3.9],
)
add_body(doc,
    "All records are loaded once at first request, parsed into a normalised list of "
    "dicts, and cached in the module-level _WORKHOUSE_CACHE variable for the lifetime "
    "of the process. Dates are parsed from free-text strings (e.g. '14 March 1849') "
    "using a regex that extracts the four-digit 19th-century year."
)

add_heading(doc, "1.3  Matching Algorithm", level=2)
add_body(doc,
    "Matching is performed by get_match_index(), which builds a dict keyed by unified "
    "record_id. The index is computed once and cached. The algorithm follows a "
    "place-first strategy with four sequential stages:"
)

add_heading(doc, "Stage 1 — Place candidate selection", level=3)
add_body(doc,
    "For each unified record, the normalised townland and parish are matched against "
    "the workhouse record's electoral_division field using three criteria in order:"
)
add_bullet(doc, "Exact equality: ed == townland or ed == parish")
add_bullet(doc, "Substring containment: ed in townland, townland in ed, ed in parish, parish in ed")
add_bullet(doc, "A pre-built inverted index (wh_by_place) enables O(1) lookup for exact matches; partial containment scans the remaining records.")
doc.add_paragraph()

add_heading(doc, "Stage 2 — Date-window filtering", level=3)
add_body(doc,
    "Within the place candidates, records are further filtered to those whose admission "
    "year falls within ±1 year of the unified record's year. This tolerates transcription "
    "errors and records where a person was admitted shortly before or after the estate "
    "event. If no place+date candidates exist, the algorithm falls back to place candidates "
    "only; if no place candidates exist, the full workhouse register is used as the scoring pool."
)

add_heading(doc, "Stage 3 — Name scoring", level=3)
add_body(doc,
    "Name similarity is computed using Python's difflib.SequenceMatcher across all "
    "cross-products of name variants. For each record, up to three variants are generated:"
)
add_bullet(doc, '"forename surname" (natural order)')
add_bullet(doc, '"surname forename" (register order — workhouse entries are stored as Surname Forename)')
add_bullet(doc, 'The raw canonical_name field from the unified record (if present)')
add_body(doc,
    "The highest SequenceMatcher ratio across all variant pairs is taken as the name "
    "score. This handles common historical issues such as inverted name order, "
    "abbreviations (Pat vs Patrick), and minor spelling variation."
)
add_code(doc, "best = max(SequenceMatcher(None, uv, wv).ratio()\n         for uv in unified_variants for wv in wh_variants)")

add_heading(doc, "Stage 4 — Occupation bonus", level=3)
add_body(doc,
    "A fixed bonus of +0.05 is added to the effective score when the unified record's "
    "occupation field and the workhouse employment field share a keyword from a curated "
    "vocabulary (labourer, farmer, servant, weaver, tailor, carpenter, etc.). The bonus "
    "improves ranking within a confidence tier but cannot promote a record across a tier "
    "boundary — it is capped at 1.0 and does not raise the raw name_score threshold."
)

add_heading(doc, "1.4  Confidence Bands", level=2)
add_body(doc,
    "Each match is assigned one of three confidence bands based on the combination of "
    "place match, date match, and name score:"
)
add_table(doc,
    headers=["Band", "Condition", "Typical use"],
    rows=[
        ["High",   "place_match AND date_match AND name_score ≥ 0.80",         "Strong evidence of the same person; suitable for direct display"],
        ["Medium", "(place_match OR date_match) AND name_score ≥ 0.60",         "Plausible match; display with caution note"],
        ["Low",    "name_score ≥ 0.60, no confirmed place or date",             "Speculative; listed only when no better match exists"],
        ["—",      "name_score < 0.60",                                          "Excluded from results entirely"],
    ],
    col_widths=[1.0, 3.2, 2.9],
)

add_heading(doc, "1.5  API Surface", level=2)
add_body(doc,
    "The matching index is exposed through a single HTTP endpoint registered as a legacy "
    "compatibility route in create_app.py:"
)
add_code(doc, "GET /api/workhouse/match/<record_id>")
add_body(doc,
    "The response is a JSON object with record_id, count, and a ranked matches array. "
    "Each match includes the raw workhouse fields, plus computed fields: confidence, "
    "name_score, location_match, occupation_match, and match_basis (a human-readable "
    "summary of what evidence contributed to the match)."
)

add_divider(doc)

add_heading(doc, "Feature 2: Application Performance", level=1)

add_heading(doc, "2.1  Scope", level=2)
add_body(doc,
    "Performance optimisation targets three layers: the SQLite database engine "
    "(PRAGMA tuning and composite indexes), in-process module-level caches for "
    "expensive or frequently repeated computations, and HTTP-level static file caching. "
    "All changes are backward-compatible with the existing raw sqlite3 architecture and "
    "introduce no new runtime dependencies."
)

add_heading(doc, "2.2  SQLite PRAGMA Tuning", level=2)
add_body(doc,
    "Five SQLite PRAGMAs are applied to every connection returned by get_db_conn() in "
    "extensions.py. Because WAL mode is enabled, these settings provide significant "
    "read throughput improvements with no durability trade-off:"
)
add_table(doc,
    headers=["PRAGMA", "Value", "Effect"],
    rows=[
        ["journal_mode", "WAL",       "Write-Ahead Logging enables concurrent reads during writes; eliminates shared-lock contention"],
        ["foreign_keys", "ON",        "Enforces referential integrity across townland / census / clearances tables"],
        ["synchronous",  "NORMAL",    "Reduces fsync calls; safe under WAL because the WAL header is always synced"],
        ["cache_size",   "-65536",    "64 MB in-memory page cache (negative value = kibibytes); reduces disk I/O for repeated queries"],
        ["temp_store",   "2",         "Temporary tables and sort results are kept in memory rather than written to a temp file"],
        ["mmap_size",    "268435456", "Maps up to 256 MB of the database file into virtual memory for lock-free read access"],
    ],
    col_widths=[1.4, 1.4, 4.3],
)

add_heading(doc, "2.3  Composite Indexes", level=2)
add_body(doc,
    "Eight covering indexes are created by ensure_schema() in extensions.py on every "
    "startup (CREATE INDEX IF NOT EXISTS is idempotent). The indexes target the query "
    "patterns generated by the census dashboard, map page, and analytics modules:"
)
add_table(doc,
    headers=["Index", "Table", "Columns", "Optimises"],
    rows=[
        ["idx_townland_civil_parish", "townland",         "civil_parish",        "Parish filter on map / census"],
        ["idx_townland_barony",       "townland",         "barony",              "Barony filter on analytics"],
        ["idx_townland_county",       "townland",         "county",              "County filter on map"],
        ["idx_townland_kg_uri",       "townland",         "kg_uri",              "KG URI lookup during ingest"],
        ["idx_census_year",           "census_record",    "year",                "Year filter on census dashboard"],
        ["idx_census_townland_year",  "census_record",    "townland_id, year",   "Townland × year range scans"],
        ["idx_clearances_year",       "clearances_record","year",                "Eviction year aggregation"],
        ["idx_clearances_tl_year",    "clearances_record","townland_id, year",   "Per-townland eviction history"],
    ],
    col_widths=[1.9, 1.5, 1.6, 2.1],
)

add_heading(doc, "2.4  In-Process Thread-Safe Caches", level=2)
add_body(doc,
    "Seven module-level caches in ask_service.py and related services avoid repeated "
    "expensive operations within a single gunicorn worker process. Each cache is "
    "protected by a threading.Lock to prevent race conditions when multiple requests "
    "arrive concurrently:"
)
add_table(doc,
    headers=["Cache", "Location", "TTL / Lifetime", "What it stores"],
    rows=[
        ["_VRTI_PARISH_CACHE",       "ask_service.py",       "1 hour (TTL)",          "SPARQL results from the public VRTI SPARQL endpoint — avoids repeated HTTP calls for the same parish/townland"],
        ["_PROMPT_SCHEMA_CACHE",     "ask_service.py",       "5 minutes (TTL)",        "The SQL schema string passed to the LLM — schema does not change at runtime"],
        ["_QUERY_MEMORY_CACHE",      "ask_service.py",       "60 seconds (TTL)",       "Approved SQL templates from the ask_query_memory table — refreshed on thumbs-up feedback"],
        ["_OPENROUTER_STATUS_CACHE", "ask_service.py",       "60 seconds (TTL)",       "OpenRouter provider health check result — avoids status ping on every question"],
        ["_OLLAMA_MODEL_CACHE",      "ask_service.py",       "2 minutes (TTL)",        "List of locally available Ollama models"],
        ["_SCHEMA_COMPAT_CACHE",     "ask_service.py",       "Process lifetime",       "Column existence check (clearances count column) that only needs to run once"],
        ["_TOWNLAND_CATALOG_CACHE",  "ask_service.py",       "Process lifetime",       "Full townland name list for fuzzy matching in the NL→SQL pipeline"],
        ["_WORKHOUSE_CACHE",         "workhouse_service.py", "Process lifetime",       "Parsed workhouse Excel rows — all two sheets loaded once into memory"],
        ["_WORKHOUSE_MATCH_INDEX",   "workhouse_service.py", "Process lifetime",       "Pre-computed record_id → ranked matches dict covering all unified records"],
        ["_GRAPH_CACHE",             "kg_service.py",        "Process lifetime",       "D3.js node/edge graph built from SQLite — expensive JOIN, cached on first call"],
    ],
    col_widths=[1.9, 1.6, 1.4, 2.2],
)

add_heading(doc, "2.5  Static File Caching", level=2)
add_body(doc,
    "The Flask application is configured with SEND_FILE_MAX_AGE_DEFAULT = 86400 "
    "(24 hours). This sets Cache-Control: max-age=86400 on all files served from "
    "frontend/static/. The two largest assets — townlands.json (~6.2 MB, Wicklow "
    "boundary polygons) and unified_processed.csv (~4.4 MB) — are served once per "
    "browser session instead of on every page load. This is particularly significant "
    "for the map page, which would otherwise re-download the full GeoJSON on every "
    "navigation."
)

add_divider(doc)

add_heading(doc, "Feature 3: RDF/KG Comparison", level=1)

add_heading(doc, "3.1  Purpose", level=2)
add_body(doc,
    "The RDF/KG Comparison feature provides a side-by-side evaluation of two query "
    "paradigms — SQL over the local SQLite database and SPARQL over an RDF knowledge "
    "graph built from the same Coolattin estate data. This supports the dissertation's "
    "comparative analysis of relational and graph-based knowledge representation, "
    "examining where the two paradigms agree, where they diverge, and why."
)

add_heading(doc, "3.2  Architecture", level=2)
add_body(doc,
    "The feature comprises two orthogonal views served from the /api/kg/* blueprint "
    "(backend/routes/kg_explore.py):"
)

add_table(doc,
    headers=["View", "Endpoint", "Description"],
    rows=[
        ["Graph topology",      "GET /api/kg/graph",     "D3.js force simulation data: nodes (Townland, Parish, EventHub, DataHub) + edges derived from the SQLite unified_record table"],
        ["SPARQL vs SQL",       "POST /api/kg/compare",  "Run a canned or custom scenario: executes the SQL in SQLite and the SPARQL via rdflib, returns both result sets side by side"],
        ["Scenario list",       "GET /api/kg/scenarios", "Returns the four canned comparison scenarios with labels and descriptions"],
        ["RDF health check",    "GET /api/kg/rdf-status","Reports whether the TTL file is present and how many triples are loaded"],
    ],
    col_widths=[1.4, 2.2, 3.5],
)

add_heading(doc, "3.3  RDF Data Source", level=2)
add_body(doc,
    "The Coolattin estate data is uplifted to RDF Turtle format by the rdf_uplift.py "
    "script. The resulting coolattin_sample.ttl file (~143,000 triples) uses a custom "
    "ontology with three namespace prefixes:"
)
add_bullet(doc, "co: (https://coolattin.ie/ontology#) — domain classes and properties: co:Person, co:Event, co:townland, co:eventType, co:year")
add_bullet(doc, "schema: (https://schema.org/) — personal name properties: schema:givenName, schema:familyName")
add_bullet(doc, "ex: (https://coolattin.ie/resource/) — individual resource URIs")
doc.add_paragraph()
add_body(doc,
    "The TTL file is loaded lazily into an rdflib.ConjunctiveGraph on first use and "
    "cached in the module-level _RDF_GRAPH variable. Loading takes approximately "
    "200–400 ms and is performed once per worker process. If the TTL file is absent "
    "(e.g. on Azure where it is gitignored), all SPARQL queries return a graceful error "
    "message rather than crashing."
)

add_heading(doc, "3.4  Canned Comparison Scenarios", level=2)
add_body(doc,
    "Four scenarios are built into kg_service.py. Each pairs an SQL query against "
    "unified_record with a semantically equivalent SPARQL query against the RDF graph, "
    "and includes a description explaining any expected difference in row counts:"
)
add_table(doc,
    headers=["Scenario ID", "Label", "Key difference"],
    rows=[
        ["emigration_count_by_townland", "Emigration by townland",
         "SQL WHERE townland IS NOT NULL matches SPARQL's open-world behaviour — triples cannot exist for NULL values"],
        ["eviction_count_by_year",       "Evictions per year",
         "SQL excludes 4 NULL-year rows to align with SPARQL, which cannot bind an absent co:year triple"],
        ["surname_frequency",            "Top 10 surnames",
         "Both return identical results — the cleanest scenario demonstrating full agreement"],
        ["person_event_detail",          "Person + event detail",
         "SPARQL traverses Person→Event as a natural graph walk; SQL implements the same join as a CASE expression in a single flat table"],
    ],
    col_widths=[1.9, 1.6, 3.6],
)

add_heading(doc, "3.5  SQL and SPARQL Runners", level=2)
add_body(doc,
    "Both query runners are implemented in kg_service.py and share the same interface "
    "(columns: list[str], rows: list[dict], error: str | None):"
)

add_heading(doc, "SQL runner (run_sql)", level=3)
add_body(doc,
    "Executes SELECT-only queries against the SQLite database via extensions.get_db_conn(). "
    "Non-SELECT statements are rejected before execution. Results are fetched with "
    "fetchmany(500) to prevent memory exhaustion on large result sets."
)
add_code(doc, "if not sql.strip().upper().startswith('SELECT'):\n    return [], [], 'Only SELECT queries are permitted.'")

add_heading(doc, "SPARQL runner (run_sparql)", level=3)
add_body(doc,
    "Executes SELECT queries against the in-process rdflib graph. Standard prefixes "
    "(co:, schema:, ex:, rdf:, rdfs:, xsd:) are prepended automatically so that "
    "scenario queries do not need to repeat them. Results are normalised: each binding "
    "value is converted to a string, with None for unbound variables."
)

add_heading(doc, "3.6  LLM Mismatch Analysis", level=2)
add_body(doc,
    "When the SQL and SPARQL row counts differ, the front end offers an 'Explain "
    "difference' action. This calls explain_mismatch() in kg_service.py, which "
    "constructs a structured academic prompt and sends it to OpenRouter (GPT / Llama / "
    "Gemma free tier, with automatic model fallback). The prompt includes both full "
    "queries and five-row samples from each result set. The LLM is instructed to "
    "respond with a fixed markdown structure:"
)
add_bullet(doc, "Root Cause — one sentence")
add_bullet(doc, "Technical Explanation — closed-world vs open-world, NULL handling, schema modelling")
add_bullet(doc, "All Possible Reasons — numbered list, most to least likely")
add_bullet(doc, "Data Evidence — what the sample rows reveal")
add_bullet(doc, "Conclusion — expected behaviour or data quality issue")
doc.add_paragraph()
add_body(doc,
    "The analysis text is returned to the front end as markdown and rendered inline "
    "alongside the comparison table. The model used is included in the response for "
    "reproducibility."
)

add_heading(doc, "3.7  Graph Topology View", level=2)
add_body(doc,
    "The /api/kg/graph endpoint builds a D3.js force simulation graph from the SQLite "
    "unified_record table. The graph uses four node types:"
)
add_table(doc,
    headers=["Node type", "ID prefix", "Colour", "What it represents"],
    rows=[
        ["Townland",  "t_<name>",         "Green / Red / Blue (dominant event)", "One node per distinct townland; size scales with total record count"],
        ["Parish",    "parish_<name>",     "Purple #7c3aed",                      "Civil parish grouping; connected to its townlands by 'contains' edges"],
        ["EventHub",  "eh_emigration etc.","Green / Red / Blue",                  "Three fixed hub nodes representing the event type taxonomy"],
        ["DataHub",   "dh_census etc.",    "Teal / Purple",                        "Two fixed hubs indicating which townlands have census or clearances data"],
    ],
    col_widths=[1.2, 1.8, 2.0, 2.1],
)
add_body(doc,
    "The graph is built by _build_graph_inner(), which executes three SQLite queries "
    "(aggregated stats per townland, parish lookup, and townland metadata) and two "
    "membership queries (which townland IDs appear in census_record and clearances_record). "
    "The result is cached in _GRAPH_CACHE on first call and served from memory on "
    "subsequent requests."
)

add_divider(doc)

add_heading(doc, "Summary of Design Decisions", level=1)
add_table(doc,
    headers=["Feature", "Key decision", "Rationale"],
    rows=[
        ["Workhouse Matching",
         "Place-first cascade with SequenceMatcher",
         "Sparse, inconsistent data means pure name-matching produces too many false positives. Restricting the scoring pool to geographically relevant records first dramatically reduces noise."],
        ["Workhouse Matching",
         "Occupation bonus capped within tier",
         "Prevents a shared occupation keyword from artificially promoting a low-quality name match. The tier boundary is a hard quality gate."],
        ["Application Performance",
         "SQLite PRAGMAs per-connection, not per-database",
         "The Flask app uses short-lived connections (no connection pool). Applying PRAGMAs in get_db_conn() ensures every connection is tuned without requiring a persistent connection."],
        ["Application Performance",
         "Process-lifetime caches for stable data",
         "The workhouse Excel, match index, and D3 graph do not change without a code or data deployment. Rebuilding them on every request would be orders of magnitude slower."],
        ["Application Performance",
         "TTL caches for external data",
         "VRTI SPARQL results, LLM provider status, and approved query memory are volatile or expensive. TTLs balance freshness against latency."],
        ["RDF/KG Comparison",
         "rdflib in-process rather than a live GraphDB connection",
         "Eliminates network dependency on the GraphDB Azure VM for the core comparison feature. The TTL file is the ground truth for the dissertation comparison."],
        ["RDF/KG Comparison",
         "LLM mismatch analysis is on-demand",
         "The LLM call is expensive (~3–8 s). It is only triggered when results actually differ and the user explicitly requests an explanation."],
    ],
    col_widths=[1.5, 2.2, 3.4],
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = p.add_run(f"Generated {datetime.datetime.now().strftime('%d %B %Y, %H:%M')}  ·  Coolattin Estate Records Explorer")
fr.font.size = Pt(9)
fr.font.color.rgb = GREY_TEXT
fr.font.italic = True

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print(f"Saved → {OUT_PATH}")
